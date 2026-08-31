"""Replace the canonical architecture figure in Blueprint submission DOCX files.

The script targets the first inline image after the architecture heading, keeps
the document's established image width, and recalculates height from the new
source asset so Word and Google Docs do not stretch it.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shape import InlineShape
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
FIGURE = ROOT / "docs" / "figures" / "architecture" / "02-end-to-end-system-architecture.png"
TARGETS = (
    (
        ROOT / "docs" / "submission" / "Blueprint-Week-3-Submission.docx",
        "End-to-End Architecture",
    ),
    (
        ROOT / "docs" / "Blueprint Evidence Dev - Final Project Documentation.docx",
        "System architecture",
    ),
)


def replace_figure(document_path: Path, heading_text: str) -> None:
    document = Document(document_path)
    heading_index = next(
        index
        for index, paragraph in enumerate(document.paragraphs)
        if heading_text.casefold() in paragraph.text.casefold()
    )

    target_paragraph = None
    for paragraph in document.paragraphs[heading_index + 1 :]:
        if paragraph._p.xpath(".//a:blip"):
            target_paragraph = paragraph
            break
        if paragraph.style.name.startswith("Heading 1") and paragraph.text.strip():
            break
    if target_paragraph is None:
        raise RuntimeError(f"No architecture image found after {heading_text!r} in {document_path}")

    blip = target_paragraph._p.xpath(".//a:blip")[0]
    relationship_id = blip.get(
        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
    )
    image_part = document.part.related_parts[relationship_id]
    image_part._blob = FIGURE.read_bytes()

    inline = target_paragraph._p.xpath(".//wp:inline")[0]
    shape = InlineShape(inline)
    with Image.open(FIGURE) as source:
        shape.height = round(shape.width * source.height / source.width)

    document.save(document_path)
    print(f"Updated {document_path.name}")


def main() -> None:
    for target, heading in TARGETS:
        replace_figure(target, heading)


if __name__ == "__main__":
    main()
